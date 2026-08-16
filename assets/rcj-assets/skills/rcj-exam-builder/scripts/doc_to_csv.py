#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
doc_to_csv.py — 原始 PDF / Word(.docx) 题库 → 结构化草稿 CSV
============================================================
把采集到的「文本型」PDF/Word 笔试题库，自动切题成草稿 CSV，
后续交给 ingest.py → build_data.py → anki_export.py 处理。

支持：
  - PDF（文本型，pypdf 抽取；扫描版会提示先 OCR）
  - Word .docx（python-docx 抽取；.doc 请另存 .docx）
  - 纯文本 .txt / .md（直接读，便于先人工粘贴）

模式：
  --type written   笔试（单选/多选/判断/填空/简答），输出列：
                   题号,批次,题型,题干,选项A,选项B,选项C,选项D,答案,解析,标签
  --type interview 面试（开放题），输出列：
                   年份,场次,题型,标题,答案,标签（答案/年份/场次尽力提取，缺失留空待填）

注意：自动切题是启发式「尽力而为」，双栏/表格题/图文混排可能不准，
      导出后请人工校对再调 skill。

依赖（缺失会提示安装）：
  pip install pypdf python-docx

用法：
  python doc_to_csv.py 题库.pdf  -o 题库_draft.csv --type written
  python doc_to_csv.py 题库.docx -o 题库_draft.csv --type interview
"""
import argparse
import csv
import os
import re
import sys

# ---------- 依赖懒加载 ----------
def _require(mod, pipname):
    try:
        return __import__(mod)
    except ImportError:
        sys.exit(
            f"[依赖缺失] 需要 {pipname}。请安装：\n"
            f"  python -m pip install {pipname}\n"
            f"然后重试。"
        )

# ---------- 文本抽取 ----------
def extract_pdf(path):
    _require("pypdf", "pypdf")
    from pypdf import PdfReader
    try:
        reader = PdfReader(path)
    except Exception as e:
        sys.exit(f"[PDF读取失败] {e}\n若是加密文件请先解密；若是扫描版请先 OCR。")
    chunks = []
    for pg in reader.pages:
        chunks.append(pg.extract_text() or "")
    full = "\n".join(chunks)
    if len(full.strip()) < 50:
        sys.exit("[疑似扫描版PDF] 抽不到足够文字，请先用 OCR 转成文本/Word，再处理。")
    return full

def extract_docx(path):
    _require("docx", "python-docx")
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

def extract_text_file(path):
    with open(path, encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".doc":
        sys.exit("[不支持 .doc] 请用 Word 另存为 .docx 再处理。")
    if ext in (".txt", ".md"):
        return extract_text_file(path)
    sys.exit(f"[不支持格式] {ext}（支持 pdf / docx / txt / md）")

# ---------- 通用正则 ----------
RE_Q = re.compile(r'^\s*(?:\(?\d{1,3}\)?|[一二三四五六七八九十百]{1,4})[\.、．]\s*\S')
RE_OPT = re.compile(r'^\s*([A-Ha-h])[\.、．]\s*(.*)')
RE_ANS = re.compile(r'(?:答案|参考答案|正确选项|答案选|标准答案)[\s:：]*([A-Ha-h](?:[\s,，、]*[A-Ha-h])*)')
RE_EXP = re.compile(r'(?:解析|详解|【解析】|【详解】|答案解析)[\s:：]*')
RE_YEAR = re.compile(r'20\d{2}')

TYPE_HINT = {
    '多选': 'multiple', '多项选择题': 'multiple', '多选题': 'multiple',
    '判断': 'judge', '判断题': 'judge',
    '填空': 'blank', '填空题': 'blank',
    '简答': 'qa', '简答题': 'qa', '问答': 'qa',
    '单选': 'single', '单项选择题': 'single', '单选题': 'single',
}

def split_blocks(text):
    lines = text.splitlines()
    blocks, cur = [], []
    for ln in lines:
        if RE_Q.match(ln):
            if cur:
                blocks.append("\n".join(cur))
            cur = [ln]
        else:
            cur.append(ln)
    if cur:
        blocks.append("\n".join(cur))
    return blocks

def guess_type(block_text, has_options):
    for kw, t in TYPE_HINT.items():
        if kw in block_text[:30]:
            return t
    if not has_options:
        return 'qa'
    return 'single'

# ---------- 笔试解析 ----------
WRITTEN_COLS = ['题号', '批次', '题型', '题干', '选项A', '选项B', '选项C', '选项D', '答案', '解析', '标签']

def parse_written(text):
    rows = []
    for blk in split_blocks(text):
        lines = [l.rstrip() for l in blk.splitlines() if l.strip()]
        if not lines:
            continue
        m = re.match(r'^\s*(?:\(?(\d{1,3})\)?|[一二三四五六七八九十百]{1,4})[\.、．]\s*(.*)', lines[0], re.S)
        if not m:
            continue
        num = m.group(1) or ""
        stem = m.group(2).strip()
        opts = {}
        answer = ""
        explanation = ""
        for l in lines[1:]:
            mo = RE_OPT.match(l)
            if mo and mo.group(2).strip():
                letter = mo.group(1).upper()
                opts[letter] = mo.group(2).strip()
                continue
            ma = RE_ANS.search(l)
            if ma:
                answer = ma.group(1).upper().replace(" ", "")
                continue
            me = RE_EXP.search(l)
            if me:
                explanation = l[me.end():].strip()
                continue
            # 续行：拼到最后一个选项，或拼到题干
            if opts:
                last = max(opts)
                opts[last] += l.strip()
            elif stem:
                stem += l.strip()
        has_opt = bool(opts)
        qtype = guess_type(blk, has_opt)
        # 判断/填空/简答不挂选项（避免 "A.对 B.错" 污染）
        if qtype in ('judge', 'blank', 'qa'):
            opts = {}
        rows.append({
            '题号': num,
            '批次': '',
            '题型': qtype,
            '题干': stem,
            '选项A': opts.get('A', ''),
            '选项B': opts.get('B', ''),
            '选项C': opts.get('C', ''),
            '选项D': opts.get('D', ''),
            '答案': answer,
            '解析': explanation,
            '标签': '',
        })
    return rows, WRITTEN_COLS

# ---------- 面试解析 ----------
INTERVIEW_COLS = ['年份', '场次', '题型', '标题', '答案', '标签']

# 常见面试题型关键词（命中即归类；都不中则标「开放题」，避免被 ingest 误归 single）
INTERVIEW_TYPE_DICT = [
    ("自我认知与职位匹配", ["自我认知", "职位匹配", "报考", "为什么", "结合自身", "优势", "不足"]),
    ("应急应变", ["应急", "突发", "遇到", "倒地", "纠纷", "冲突", "紧急", "险情"]),
    ("人际关系", ["同事", "领导", "群众", "沟通", "关系", "配合"]),
    ("组织协调", ["组织", "活动", "会议", "调研", "培训", "安排", "策划"]),
    ("综合分析", ["怎么看", "谈谈", "现象", "看法", "观点", "认为", "理解"]),
    ("演讲与情景模拟", ["演讲", "情景模拟", "现场模拟"]),
]

def guess_interview_type(title):
    for t, kws in INTERVIEW_TYPE_DICT:
        if any(k in title for k in kws):
            return t
    return "开放题"

def parse_interview(text):
    raw_lines = [l.rstrip() for l in text.splitlines()]
    cur_year = ""
    cleaned = []  # (year, line) —— 先剥离「年份标题行」，避免它污染题目
    for ln in raw_lines:
        s = ln.strip()
        if not s:
            continue
        ym = RE_YEAR.search(s)
        if ym and (len(s) < 30 or any(k in s for k in ("真题", "面试", "试题", "招考"))):
            cur_year = ym.group(0)   # 仅作年份上下文，不单独成题
            continue
        cleaned.append((cur_year, s))
    # 仅按题号分块
    blocks, cur = [], []
    for year, ln in cleaned:
        if RE_Q.match(ln):
            if cur:
                blocks.append(cur)
            cur = [(year, ln)]
        else:
            cur.append((year, ln))
    if cur:
        blocks.append(cur)
    rows = []
    for blk in blocks:
        blk_year = blk[0][0]
        qlines = [l for _, l in blk]
        m = re.match(r'^\s*(?:\(?(\d{1,3})\)?|[一二三四五六七八九十百]{1,4})[\.、．]\s*(.*)', qlines[0], re.S)
        title = (m.group(2).strip() if m else qlines[0].strip())
        for l in qlines[1:]:
            if RE_OPT.match(l) or RE_ANS.search(l):
                break
            title += l.strip()
        rows.append({
            '年份': blk_year,
            '场次': '',
            '题型': guess_interview_type(title),
            '标题': title,
            '答案': '',
            '标签': '',
        })
    return rows, INTERVIEW_COLS

def write_csv(rows, cols, out):
    with open(out, 'w', encoding='utf-8-sig', newline='') as f:
        w = csv.DictWriter(f, fieldnames=cols)
        w.writeheader()
        for r in rows:
            w.writerow(r)
    print(f"[+] 已写入 {out}：共 {len(rows)} 题（草稿，请校对后交 ingest.py）")

def main():
    ap = argparse.ArgumentParser(description="PDF/Word 题库 → 草稿 CSV")
    ap.add_argument("input", help="源文件 .pdf/.docx/.txt/.md")
    ap.add_argument("-o", "--out", help="输出 CSV，默认 <输入名>_draft.csv")
    ap.add_argument("--type", choices=["written", "interview"], default="written")
    args = ap.parse_args()
    out = args.out or (os.path.splitext(args.input)[0] + "_draft.csv")
    if not os.path.exists(args.input):
        sys.exit(f"[!] 文件不存在: {args.input}")
    text = extract(args.input)
    if args.type == "written":
        rows, cols = parse_written(text)
    else:
        rows, cols = parse_interview(text)
    if not rows:
        sys.exit("[!] 未识别到题目。请检查文件是否为文本型，或改用 .txt 手工粘贴后重试。")
    write_csv(rows, cols, out)

if __name__ == "__main__":
    main()
